import os
import sys
import time
import operator
import logging
import requests
import json
import re
import docx
import subprocess
import codecs

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from watchdog.events import FileModifiedEvent
from watchdog.events import FileCreatedEvent

from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfdocument import PDFNoOutlines

from nltk.stem import WordNetLemmatizer


class GamificationHandler(FileSystemEventHandler):

    def __init__(self, paper_filename, publish_url, paper_id):
        FileSystemEventHandler.__init__(self)  # super init

        self.paper_filename = paper_filename
        self.is_pdf = paper_filename.endswith('pdf')

        self.publish_url = publish_url
        self.paper_id = paper_id

        self.reset_stats()

        logging.info(
            "Creating a GamificationHandler with paper: " + paper_filename +
            " publish_url: " + publish_url +
            " and paper id: " + paper_id
        )

    def reset_stats(self):
        self.stats = {}
        self.words = {}
        self.paragraphs = []
        self.num_words = 0
        self.total_word_len = 0
        self.pages = None

    def on_created(self, event):
        # MAIN CALLBACK - a file got created
        logging.info("Create event occurred: " + event.src_path)
        if type(event) == FileCreatedEvent:
            logging.info("A file was created: " + event.src_path)

            self.analyze_file_event(event)

    def on_modified(self, event):
        # MAIN CALLBACK - a file got modified
        logging.info("Modify event occurred: " + event.src_path)
        if type(event) == FileModifiedEvent:
            logging.info("A file was modified: " + event.src_path)

            self.analyze_file_event(event)

    def analyze_file_event(self, event):
        paper_path = os.path.abspath(self.paper_filename)
        logging.info("Checking if it was the paper: " + paper_path)

        if paper_path == event.src_path:
            logging.info("Paper change detected, calculating statistics ...")
            self.analyze_paper()

    def analyze_paper(self):
        self.reset_stats()
        self.parse_file()
        self.calculate_statistics()
        logging.info("Publishing ...")
        self.publish()
        logging.info("Published!")

    def parse_paragraphs(self, text):
        # Will only work for markdown elements
        #   divided by '##' markers
        #   or for pdf like chapters, e.g. \n\n 2 Conclusion \n\n
        lines = text.split('\n')
        headlines = []

        if self.is_pdf:
            with open(self.paper_filename, 'rb') as pdf:
                parser = PDFParser(pdf)
                document = PDFDocument(parser)

                try:
                    outlines = document.get_outlines()
                    for (level, title, _, _, _) in outlines:
                        if level == 1:
                            headlines.append(title)
                except PDFNoOutlines:
                    logging.info(
                        "No outline found -> skipping paragraph search..."
                    )
        else:  # check markdown headlines
            for index, line in enumerate(lines):
                if line.startswith('## '):
                    headlines.append(line)

        if len(headlines) > 0:
            self.count_paragraphs(text, lines, headlines)

    def compress_line(self, line):
        return line.replace(' ', '').lower()

    def count_paragraphs(self, text, lines, headlines):
        old_headline = ""
        compressed_headlines = list(self.compress_line(h) for h in headlines)

        for line in lines:
            compressed_line = self.compress_line(line)

            if line in headlines or compressed_line in compressed_headlines:
                if old_headline != "":
                    # Count previous paragraph
                    paragraph = text.split(old_headline)[1].split(line)[0]
                    if self.is_pdf:
                        old_headline = headlines[
                            compressed_headlines.index(compressed_line) - 1
                        ]
                    self.count_paragraph_words(old_headline, paragraph)
                old_headline = line

        # Count last paragraph
        if old_headline != "":
            paragraph = text.split(old_headline)[1]
            if self.is_pdf:
                old_headline = headlines[
                    compressed_headlines.index(self.compress_line(old_headline)) - 1
                ]
            self.count_paragraph_words(old_headline, paragraph)

    def count_paragraph_words(self, line, paragraph):
        num_words = len(re.findall(r"[\w']+", paragraph))
        self.paragraphs.append((line.replace('#', '').strip(), num_words))

    def parse_text_statistics(self, text):
        wnl = WordNetLemmatizer()
        for w in text:
            word = wnl.lemmatize(w.strip().lower())

            # Add to total_word_len
            # to determine average word length later
            self.total_word_len += len(word)

            # Count distinct words with occurrences
            if word not in self.words:
                self.words[word] = 0
            self.words[word] += 1

            # Count all words
            self.num_words += 1

    def get_pages(self):
        fp = open(self.paper_filename, 'rb')
        self.pages = 0
        for page in PDFPage.get_pages(fp):
            self.pages += 1
        fp.close()

    def parse_pdf_file(self):
        # Convert pdf to txt
        tmp_filename = "tmpExtracted.txt"
        pdf_convert_exit_id = subprocess.call(
            ["pdf2txt.py", "-o", tmp_filename, self.paper_filename]
        )
        if pdf_convert_exit_id == 0:
            logging.info("\t\t\tSuccessfully converted pdf to txt")
            # Analyse plain text
            logging.info("\t\t\tAnalyzing file")
            text = self.analyze_file(tmp_filename)
            logging.info("\t\t\tParsing paragraphs ...")
            self.parse_paragraphs(text)
            logging.info("\t\t\tGetting pages ...")
            self.get_pages()

    def parse_word_file(self):
        # Read file
        document = docx.Document(self.paper_filename)
        text = ""
        for p in document.paragraphs:
            text += p.text + '\n'

        word_split = re.findall(r"[\w']+", text)

        # Analyse
        self.parse_text_statistics(word_split)

    def parse_text_file(self):
        text = self.analyze_file(self.paper_filename)
        self.parse_paragraphs(text)

    def analyze_file(self, filename):
        f = codecs.open(filename, "r", "utf-8")
        text = ""
        for line in f.readlines():
            text += line
            word_split = re.findall(r"[\w']+", line)
            # Analyse
            self.parse_text_statistics(word_split)

        f.close()
        return text

    def parse_file(self):
        logging.info("\tParsing the paper ...")
        if self.paper_filename.endswith(".docx"):
            logging.info("\t\tusing docx parser ...")
            self.parse_word_file()
        elif self.paper_filename.endswith(".pdf"):
            logging.info("\t\tusing pdf parser ...")
            self.parse_pdf_file()
        else:
            logging.info("\t\tusing txt parser ...")
            self.parse_text_file()

    def calculate_statistics(self):
        # By now, text-statistics should be saved in instance variables

        # Determine interesting words
        logging.info("\tCalculating interesting words ...")
        interesting_words = self.get_interesting_words(40)

        # Determine average word length
        logging.info("\tCalculating average word length ...")
        avg_len = float(self.total_word_len) / float(self.num_words)

        # Determine Oxford coverage
        logging.info("\tCalculating oxford coverage ...")
        oxford_coverage = self.get_coverage("./oxford.txt")

        # Determine Fancy word coverage
        logging.info("\tCalculating fancy words coverage ...")
        fancy_coverage = self.get_coverage("./fancy.txt")

        # Determine academic word list coverage
        logging.info("\tCalculating academic word list coverage ...")
        awl_coverage = self.get_awl_coverage("./awl.txt")

        # Build stats together
        logging.info("\tBuilding stats together ...")

        self.stats = {
            "num_words": self.num_words,
            "different_words": len(self.words),
            "avg_len": avg_len,
            "paragraphs": self.paragraphs,
            "interesting_words": interesting_words,
            "oxford_coverage": {
                "total": oxford_coverage["total"],
                "num_hits": len(oxford_coverage["hits"])
            },
            "fancy_coverage": {
                "total": fancy_coverage["total"],
                "num_hits": len(fancy_coverage["hits"])
            },
            "awl_coverage": {
                "words_total": awl_coverage["words_total"],
                "words_hits": awl_coverage["words_hits"],
                "category_total": awl_coverage["category_total"],
                "category_num_hits": awl_coverage["category_num_hits"],
                "category_hits": awl_coverage["category_hits"]
            }
        }

        if self.pages is not None:
            self.stats["pages"] = self.pages

        logging.info("\tStats: " + str(self.stats))

    def get_interesting_words(self, num):
        sorted_words = sorted(
            self.words.iteritems(), key=operator.itemgetter(1), reverse=True
        )
        interesting_words = []

        num = min(num, len(sorted_words))
        min_len = 10

        while len(interesting_words) != num:
            # As long as we don't have as many words as we want
            for word in sorted_words:
                if len(word[0]) >= min_len:
                    if word[1] == 1:
                        # Word only occurs once in the text
                        # -> since sorted_words is sorted by occurrence:
                        #    break and go down with min word length
                        break
                    if word not in interesting_words:
                        interesting_words.append(word)
                if len(interesting_words) == num:
                    # Got enough words, break will break both loops
                    break
            min_len -= 1
            if min_len < 2:
                # Text contains really few words, we just have to add them
                #  until we have enough
                for word in sorted_words:
                    if word not in interesting_words:
                        interesting_words.append(word)
                        if len(interesting_words) == num:
                            # Got enough words, break will break both loops
                            break

        # Sort result and return
        interesting_words = sorted(
            interesting_words, key=operator.itemgetter(1), reverse=True
        )
        return interesting_words

    def get_coverage(self, filename):
        """ Reads a list of words and compares it to the own words"""
        words = []
        num_words = 0
        # Count and compare
        f = open(filename)
        for word in f.readlines():
            if word.strip() != "":
                words.append(word.strip().lower())
                num_words += 1
        f.close()

        hits = set(words).intersection(set(self.words.keys()))

        return {"total": num_words, "hits": list(hits)}

    def get_awl_coverage(self, filename):
        words = {}
        f = open(filename)

        category = ""
        for word in f.readlines():
            if not word.startswith('\t'):
                category = word.strip()
            words[word.strip()] = category

        hits = set(words.keys()).intersection(set(self.words.keys()))

        category_hits = {}
        for category in set(words.values()):
            category_hits[category] = 0

        for hit in hits:
            category_hits[words[hit]] += 1

        category_num_hits = 0
        for key in category_hits.keys():
            if category_hits[key] > 0:
                category_num_hits += 1

        return {
            "words_total": len(words),
            "words_hits": len(list(hits)),
            "category_total": len(list(set(words.values()))),
            "category_num_hits": category_num_hits,
            "category_hits": category_hits
        }

    def publish(self):
        payload = {"stats": json.dumps(self.stats)}
        requests.put(
            self.publish_url + "/papers/" + self.paper_id + ".json",
            data=payload
        )


def set_paper_alive(publish_url, paper_id, alive):
    payload = {"alive": str(alive).lower()}
    requests.put(
        publish_url + "/papers/" + paper_id + ".json",
        params=payload
    )


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO,
                        format='%(asctime)s - %(message)s',
                        datefmt='%Y-%m-%d %H:%M:%S')
    if len(sys.argv) != 4:
        print "Usage: python tracker.py <paper-file> <publish-host> <paper-id>"
        sys.exit()

    # Parse command line params
    filename = sys.argv[1]
    publish_url = sys.argv[2]
    paper_id = sys.argv[3]

    path = os.path.dirname(os.path.abspath(filename))

    # Enable "Currently writing..."
    set_paper_alive(publish_url, paper_id, True)

    # Observer setup
    event_handler = GamificationHandler(filename, publish_url, paper_id)
    observer = Observer()
    logging.info("Starting observer with watch path: " + path)
    observer.schedule(event_handler, path=path, recursive=True)
    # Observer start
    observer.start()
    logging.info("Observer started.")

    # Trigger an initial paper analysis
    event_handler.analyze_paper()

    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        # Disable "Currently writing..."
        set_paper_alive(publish_url, paper_id, False)
        observer.stop()
        observer.join()
